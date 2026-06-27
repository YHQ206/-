"""Word文件题库解析器 - 解析客观题题库及答案.docx

答案标记方式：
- 单选/多选题：红色字体(FF0000)标记正确选项
- 判断题：行内(√)/(×)标记，或在表格中
- 填空题：答案在括号内
"""
import re
import docx
from docx.shared import RGBColor
from docx.oxml.ns import qn


def _is_red(run):
    """判断run是否为红色字体"""
    return run.font.color and run.font.color.rgb == RGBColor(0xFF, 0x00, 0x00)


def _has_underline(run):
    """判断run是否有下划线"""
    if run.font.underline:
        return True
    # 直接检查XML（有些docx的underline信息在XML里但python-docx读不到）
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        u = rPr.find(qn('w:u'))
        if u is not None:
            val = u.get(qn('w:val'))
            if val is None or val == 'single':
                return True
    return False


def _get_underline_texts(para):
    """获取段落中所有下划线标记的文本片段列表"""
    texts = []
    for run in para.runs:
        if _has_underline(run) and run.text.strip():
            texts.append(run.text.strip())
    return texts


def _extract_chapter_names(doc):
    """从目录中提取章节名称"""
    chapters = []
    for para in doc.paragraphs:
        text = para.text.strip()
        m = re.match(r'第\s*(\d+)\s*章\s+(.+?)(?:\t|$)', text)
        if m:
            chapters.append(f'第{m.group(1)}章 {m.group(2).strip()}')
    return chapters


def _extract_chapter_names_from_toc(doc):
    """从单行目录中提取章节名称（适配学习通格式）

    目录格式示例：
    导论xxx...2 第 1 章xxx...9 第 2 章xxx...15 第 4 章xxx...20 ...
    """
    # 查找包含多个章节标记的目录行
    for para in doc.paragraphs:
        text = para.text.strip()
        # 包含"导论"和"第X章"的行是目录
        if '导论' in text and '第' in text and '章' in text:
            chapters = []
            # 提取导论（导论后面是章节名，然后是点和页码）
            m = re.search(r'导论(.+?)\.{2,}', text)
            if m:
                chapters.append('导论 ' + m.group(1).strip())
            # 提取各章（第 X 章 + 章节名 + 点和页码）
            for m in re.finditer(r'第\s*(\d+)\s*章\s*(.+?)\.{2,}', text):
                ch_name = f'第{m.group(1)}章 {m.group(2).strip()}'
                chapters.append(ch_name)
            if chapters:
                return chapters
    # 备用：尝试每行一个章节名的格式
    return _extract_chapter_names(doc)


def _get_line_red_answer_keys(para):
    """获取一行中红色标记的选项字母列表

    返回: (keys, has_red, red_full_text)
        keys: 红色标记的选项字母列表
        has_red: 是否有红色文本
        red_full_text: 红色文本的完整内容（用于整行红色的情况）
    """
    # 拼接所有红色run的文本
    red_text = ''
    for run in para.runs:
        if _is_red(run):
            red_text += run.text

    if not red_text:
        return [], False, ''

    # 匹配 （A） 格式
    m = re.findall(r'（([A-H])）', red_text)
    if m:
        return m, True, red_text

    # 匹配独立的字母
    m = re.findall(r'(?<![a-zA-Za-z])([A-H])(?![a-zA-Za-z])', red_text)
    if m:
        return m, True, red_text

    # 有红色文本但没有匹配到字母 - 返回完整文本用于后续处理
    return [], True, red_text


def _parse_options_from_text(text):
    """从文本中解析选项，返回 [(key, text), ...]"""
    options = []

    # 先尝试用正则直接匹配所有 （X）选项 格式
    # 这样可以处理tab分隔、空格分隔等各种情况
    pattern = r'（([A-H])）\s*(.+?)(?=（[A-H]）|$)'
    matches = re.findall(pattern, text)
    if matches:
        for key, content in matches:
            # 清理内容中的制表符和多余空格
            content = re.sub(r'\s+', ' ', content).strip()
            options.append((key, content))
        # 检查第一个选项前面是否有无标记的文本
        first_match = re.search(r'（([A-H])）', text)
        if first_match and first_match.start() > 0:
            prefix = text[:first_match.start()].strip()
            if prefix:
                options.insert(0, ('', prefix))
        return options

    # 按tab分割（备用方案）
    parts = text.split('\t')
    for part in parts:
        part = part.strip()
        if not part:
            continue
        m = re.match(r'（([A-H])）\s*(.+)', part)
        if m:
            options.append((m.group(1), m.group(2).strip()))
            continue
        # 纯文本选项（没有标记）
        if part and not re.match(r'^[A-H]$', part):
            options.append(('', part))

    return options


def _parse_options_with_first_unmarked(text):
    """解析选项，处理第一个选项没有（A）标记的情况

    例如：十七大	（B）十八大	（C）十九大	（D）二十大
    应该解析为：A=十七大, B=十八大, C=十九大, D=二十大
    也支持空格分隔：（A）XXX （B）XXX （C）XXX
    """
    # 先用正则匹配所有 （X）内容 格式（支持tab和空格分隔）
    pattern = r'（([A-H])）\s*(.+?)(?=\s*（[A-H]）|$)'
    matches = re.findall(pattern, text)
    if matches and len(matches) >= 2:
        options = []
        for key, content in matches:
            content = re.sub(r'\s+', ' ', content).strip()
            options.append((key, content))
        # 检查第一个选项前面是否有无标记的文本
        first_match = re.search(r'（([A-H])）', text)
        if first_match and first_match.start() > 0:
            prefix = text[:first_match.start()].strip()
            if prefix:
                options.insert(0, ('', prefix))
        # 如果有无标记的选项，分配字母
        if any(k == '' for k, _ in options):
            options = _assign_option_keys(options)
        return options

    # 正则没匹配到（或只有1个匹配），按tab分割（兼容旧格式）
    options = []
    parts = text.split('\t')

    for part in parts:
        part = part.strip()
        if not part:
            continue
        m = re.match(r'（([A-H])）\s*(.+)', part)
        if m:
            options.append((m.group(1), m.group(2).strip()))
        else:
            options.append(('', part))

    # 如果有无标记的选项，给它们分配字母
    if options and any(k == '' for k, _ in options):
        options = _assign_option_keys(options)

    return options


def _assign_option_keys(options):
    """给无标记的选项分配字母"""
    used_keys = {k for k, _ in options if k}
    new_options = []
    next_key_idx = 0
    letters = 'ABCDEFGH'
    for key, text in options:
        if key == '':
            while next_key_idx < len(letters) and letters[next_key_idx] in used_keys:
                next_key_idx += 1
            if next_key_idx < len(letters):
                key = letters[next_key_idx]
                used_keys.add(key)
                next_key_idx += 1
        new_options.append((key, text))
    return new_options


def _is_option_line(text):
    """判断是否是选项行"""
    # 包含 （B）, （C）, （D） 等标记
    if re.search(r'（[B-H]）', text):
        return True
    # 包含 （A） 标记
    if '（A）' in text:
        return True
    return False


def _is_new_stem_line(text):
    """判断是否是新题干的开始（用于v2解析器）"""
    # 包含填空下划线
    if '________' in text or '______' in text:
        return True
    # 以句号、问号、感叹号结尾（且较长，排除短选项）
    if text.endswith(('。', '？', '！')) and len(text) > 15:
        return True
    # 以逗号结尾（题干续行信号，但不是新题干）
    # 不包含在这里，因为逗号结尾更可能是题干续行
    return False


def _add_option_to_question(q, text, has_red, red_keys):
    """给题目添加一个选项"""
    existing_keys = {o[0] for o in q['options']}
    for letter in 'ABCDEFGH':
        if letter not in existing_keys:
            q['options'].append((letter, text))
            # 红色标记的是答案
            if has_red:
                q['red_keys'].append(letter)
            elif red_keys:
                q['red_keys'].extend(red_keys)
            break


def _is_likely_question_stem(text):
    """判断是否是题干行（启发式方法）"""
    # 以问号、句号、冒号结尾的通常是题干
    if text.endswith(('？', '。', '：', '?', '.', ':')):
        return True
    # 包含"下列"、"以下"等词通常是题干开头
    if any(word in text for word in ['下列', '以下', '以下哪', '下列哪']):
        return True
    # 以逗号结尾可能是题干续行
    if text.endswith(('，', ',')):
        return True
    # 包含句号在中间，可能是题干
    if '。' in text and len(text) > 20:
        return True
    # 以"是"结尾通常是题干（用于"____是"的模式）
    if text.endswith('是') and len(text) > 5:
        return True
    # 以"其"结尾通常是题干
    if text.endswith('其') and len(text) > 5:
        return True
    # 包含问号通常是题干
    if '？' in text or '?' in text:
        return True
    # 以"是"开头且较长可能是题干（但短行可能是选项）
    if text.startswith('是') and len(text) > 15:
        return True
    # 包含"是指"通常是题干
    if '是指' in text:
        return True
    return False


def parse_docx_file(filepath):
    """解析Word格式的题库文件，返回题目列表"""
    doc = docx.Document(filepath)

    # 提取章节名称
    chapter_names = _extract_chapter_names(doc)

    # 收集表格中的判断题答案
    table_judge_answers = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2:
                q_text = cells[0]
                ans_text = cells[1]
                if '（√）' in ans_text or '（×）' in ans_text:
                    answer = '√' if '√' in ans_text else '×'
                    table_judge_answers[q_text[:20]] = answer

    # 解析所有段落
    all_items = []
    current_type = None
    chapter_idx = 0

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if para_idx < 28:  # 跳过目录
            continue

        # 检测题型切换
        if text == '一、单选题':
            current_type = 'single'
            chapter_idx = _detect_chapter_index(para_idx, doc)
            continue
        elif text == '二、多选题':
            current_type = 'multiple'
            continue
        elif text == '三、填空题':
            current_type = 'blank'
            continue
        elif text == '四、判断题':
            current_type = 'judge'
            continue
        elif text in ('五、简答题', '六、材料论述题'):
            current_type = None
            continue

        if current_type is None:
            continue

        # 跳过答案行
        if text.startswith('【答案】') or text.startswith('【参考答案】'):
            continue

        # 获取红色标记的选项
        red_keys, has_red, red_full_text = _get_line_red_answer_keys(para)
        # 获取下划线文本（填空题答案用）
        underline_texts = _get_underline_texts(para)

        all_items.append((para_idx, current_type, text, red_keys, has_red, chapter_idx, red_full_text, underline_texts))

    # 解析所有条目
    questions = []
    _parse_all_items(all_items, questions, table_judge_answers, chapter_names)

    return questions


def _detect_chapter_index(para_idx, doc):
    """检测当前段落属于哪个章节"""
    count = 0
    for i in range(para_idx + 1):
        text = doc.paragraphs[i].text.strip()
        if text == '一、单选题':
            count += 1
    return count - 1


def _parse_all_items(all_items, questions, table_judge_answers, chapter_names):
    """解析所有条目，生成题目列表"""
    # 按题型分组
    groups = []
    current_type = None
    current_items = []
    current_chapter_idx = 0

    for item in all_items:
        para_idx, qtype, text, red_keys, has_red, chapter_idx, red_full_text, underline_texts = item

        if qtype != current_type:
            if current_items:
                groups.append((current_type, current_items[:], current_chapter_idx))
            current_type = qtype
            current_items = []
            current_chapter_idx = chapter_idx

        current_items.append((para_idx, text, red_keys, has_red, red_full_text, underline_texts))

    if current_items:
        groups.append((current_type, current_items[:], current_chapter_idx))

    # 解析每组
    for qtype, items, chapter_idx in groups:
        chapter_name = chapter_names[chapter_idx] if chapter_idx < len(chapter_names) else f'第{chapter_idx+1}章'

        if qtype == 'single':
            _parse_single_choice(items, questions, chapter_name)
        elif qtype == 'multiple':
            _parse_multiple_choice(items, questions, chapter_name)
        elif qtype == 'blank':
            _parse_blanks(items, questions, chapter_name)
        elif qtype == 'judge':
            _parse_judges(items, questions, chapter_name, table_judge_answers)


def _parse_single_choice(items, questions, chapter_name):
    """解析单选题"""
    current_q = None

    i = 0
    while i < len(items):
        para_idx, text, red_keys, has_red, red_full_text, _ul = items[i]

        if _is_option_line(text):
            # 显式选项行
            opts = _parse_options_with_first_unmarked(text)
            if current_q is not None:
                # 处理选项
                for key, opt_text in opts:
                    current_q['options'].append((key, opt_text))
                    # 如果有红色文本，检查是否包含这个key
                    if red_keys and key in red_keys:
                        current_q['red_keys'].append(key)
                    elif has_red and red_full_text:
                        # 检查红色文本是否包含这个选项的内容
                        if opt_text in red_full_text or red_full_text in opt_text:
                            current_q['red_keys'].append(key)
        elif _is_likely_question_stem(text) or (current_q is None):
            # 题干行 - 检查是否需要保存之前的题目
            if current_q and current_q['stem']:
                # 检查当前题目是否已经有选项
                # 如果没有选项，可能是题干跨行
                if not current_q['options']:
                    # 检查当前题干是否以句号结尾
                    if current_q['stem'].endswith(('。', '？', '！')):
                        # 以句号结尾，保存之前的题目
                        _save_choice_question(current_q, 'single', questions, chapter_name)
                        current_q = {'stem': text, 'options': [], 'red_keys': []}
                    else:
                        # 不以句号结尾，可能是题干跨行
                        current_q['stem'] += text
                else:
                    # 已经有选项，保存之前的题目
                    _save_choice_question(current_q, 'single', questions, chapter_name)
                    current_q = {'stem': text, 'options': [], 'red_keys': []}
            else:
                current_q = {'stem': text, 'options': [], 'red_keys': []}

            # 检查这行是否同时包含题干和选项
            if '（B）' in text or '（C）' in text:
                stem_part = text.split('（B）')[0].split('（C）')[0].strip()
                if stem_part:
                    opts = _parse_options_with_first_unmarked(text)
                    current_q = {'stem': stem_part, 'options': opts, 'red_keys': red_keys}
                else:
                    current_q = {'stem': text, 'options': [], 'red_keys': []}
        else:
            # 可能是无标记的选项行或题干续行
            # 检查是否是短行且不是题干
            if len(text) < 30 and not text.endswith(('。', '？', '！', '，', ',')):
                # 可能是选项
                if current_q is not None:
                    # 检查是否是题干续行
                    # 如果当前题干不以句号结尾，可能是续行
                    if current_q['stem'] and not current_q['stem'].endswith(('。', '？', '！')):
                        # 检查是否已经有选项
                        # 如果当前题目已经有选项，那么这行也可能是选项
                        if current_q['options']:
                            # 已经有选项，这行也是选项
                            existing_keys = {o[0] for o in current_q['options']}
                            for letter in 'ABCDEFGH':
                                if letter not in existing_keys:
                                    current_q['options'].append((letter, text))
                                    # 如果有红色文本，记录答案
                                    if red_keys:
                                        current_q['red_keys'].extend(red_keys)
                                    elif has_red:
                                        # 整行都是红色，这个选项就是答案
                                        current_q['red_keys'].append(letter)
                                    break
                        else:
                            # 没有选项，可能是题干续行
                            current_q['stem'] += text
                    else:
                        # 根据已有选项推断key
                        existing_keys = {o[0] for o in current_q['options']}
                        for letter in 'ABCDEFGH':
                            if letter not in existing_keys:
                                current_q['options'].append((letter, text))
                                # 如果有红色文本，记录答案
                                if red_keys:
                                    current_q['red_keys'].extend(red_keys)
                                elif has_red:
                                    # 整行都是红色，这个选项就是答案
                                    current_q['red_keys'].append(letter)
                                break
                else:
                    # 没有当前题目，创建一个
                    current_q = {'stem': text, 'options': [], 'red_keys': []}
            else:
                # 题干行或题干续行
                if current_q and current_q['stem']:
                    # 检查是否是题干续行（以逗号结尾或不以句号结尾）
                    if text.endswith(('，', ',')) or (not text.endswith(('。', '？', '！')) and len(text) < 30):
                        # 题干续行
                        current_q['stem'] += text
                    else:
                        # 新题干
                        _save_choice_question(current_q, 'single', questions, chapter_name)
                        current_q = {'stem': text, 'options': [], 'red_keys': []}
                else:
                    current_q = {'stem': text, 'options': [], 'red_keys': []}

        i += 1

    # 保存最后一题
    if current_q and current_q['stem']:
        _save_choice_question(current_q, 'single', questions, chapter_name)


def _parse_multiple_choice(items, questions, chapter_name):
    """解析多选题"""
    current_q = None

    for para_idx, text, red_keys, has_red, red_full_text, _ul in items:
        if _is_option_line(text):
            opts = _parse_options_with_first_unmarked(text)
            if current_q is not None:
                for key, opt_text in opts:
                    current_q['options'].append((key, opt_text))
                    # 如果有红色文本，检查是否包含这个key
                    if red_keys and key in red_keys:
                        current_q['red_keys'].append(key)
                    elif has_red and red_full_text:
                        # 检查红色文本是否包含这个选项的内容
                        if opt_text in red_full_text or red_full_text in opt_text:
                            current_q['red_keys'].append(key)
        elif _is_likely_question_stem(text):
            # 题干行
            if current_q and current_q['stem']:
                _save_choice_question(current_q, 'multiple', questions, chapter_name)

            current_q = {'stem': text, 'options': [], 'red_keys': []}
        else:
            # 可能是无标记的选项行
            if current_q is not None:
                # 检查是否是短行（可能是选项）
                if len(text) < 50:
                    # 短行，可能是选项
                    existing_keys = {o[0] for o in current_q['options']}
                    for letter in 'ABCDEFGH':
                        if letter not in existing_keys:
                            current_q['options'].append((letter, text))
                            # 如果有红色文本，记录答案
                            if red_keys:
                                current_q['red_keys'].extend(red_keys)
                            elif has_red:
                                # 整行都是红色，这个选项就是答案
                                current_q['red_keys'].append(letter)
                            break
                else:
                    # 长行，可能是题干
                    _save_choice_question(current_q, 'multiple', questions, chapter_name)
                    current_q = {'stem': text, 'options': [], 'red_keys': []}
            else:
                current_q = {'stem': text, 'options': [], 'red_keys': []}

    # 保存最后一题
    if current_q and current_q['stem']:
        _save_choice_question(current_q, 'multiple', questions, chapter_name)


def _save_choice_question(q, qtype, questions, chapter_name):
    """保存一道选择题"""
    options = q['options']
    red_keys = q['red_keys']

    # 如果没有显式选项，可能是选项格式不规范
    if not options:
        # 检查题干是否包含选项
        stem = q['stem']
        if '（B）' in stem or '（C）' in stem:
            # 从题干中提取选项
            parts = re.split(r'（([A-H])）', stem)
            if len(parts) >= 3:
                q['stem'] = parts[0].strip()
                for i in range(1, len(parts), 2):
                    if i + 1 < len(parts):
                        key = parts[i]
                        text = parts[i + 1].strip()
                        options.append((key, text))

    # 构建选项列表
    option_list = []
    for key, text in options:
        if key:  # 有明确的key
            option_list.append({'key': key, 'text': text})
        else:
            # 没有明确的key，可能是第一个选项
            # 根据已有选项推断
            existing_keys = {o['key'] for o in option_list}
            for letter in 'ABCDEFGH':
                if letter not in existing_keys:
                    option_list.append({'key': letter, 'text': text})
                    break

    # 构建答案
    answer = ''.join(sorted(set(red_keys))) if red_keys else ''

    # 处理题干中的空白
    stem = q['stem']
    # 将制表符替换为横线
    stem = stem.replace('\t', '____')
    # 处理不完整的句子（缺少主语的情况）
    # 例如：宣布，中国特色社会主义进入了新时代
    # 应该显示为：____宣布，中国特色社会主义进入了新时代
    if stem and not stem.startswith('____'):
        # 检查是否以动词开头（缺少主语）
        if stem[0] in '宣布是坚持推动促进完善加强深化':
            stem = '____' + stem
        # 检查是否以"是"开头
        elif stem.startswith('是'):
            stem = '____' + stem

    questions.append({
        'type': qtype,
        'content': stem,
        'options': option_list if option_list else None,
        'answer': answer,
        'chapter': chapter_name
    })


def _parse_blanks(items, questions, chapter_name):
    """解析填空题"""
    for item in items:
        para_idx, text, red_keys, has_red, red_full_text, underline_texts = item
        m = re.match(r'^(\d+)[.、]\s*(.+)', text)
        if m:
            q_content = m.group(2).strip()
        else:
            q_content = text

        # 优先使用下划线文本作为答案（docx中的下划线标记填空答案）
        if underline_texts:
            answer = '、'.join(underline_texts)
            # 将下划线部分替换为____
            clean_content = q_content
            for ut in underline_texts:
                clean_content = clean_content.replace(ut, '____', 1)
            questions.append({
                'type': 'blank',
                'content': clean_content,
                'options': None,
                'answer': answer,
                'chapter': chapter_name
            })
            continue

        # 从内容中提取括号内的答案
        answers = re.findall(r'（(.+?)）', q_content)
        if not answers:
            answers = re.findall(r'\((.+?)\)', q_content)

        if answers:
            clean_content = re.sub(r'（.+?）', '____', q_content)
            clean_content = re.sub(r'\(.+?\)', '____', clean_content)
            answer = '、'.join(a.strip() for a in answers)
        else:
            # 没有括号答案，尝试从句子中推断答案
            # 使用更智能的方式识别答案
            answer, clean_content = _extract_blank_answer(q_content)

        questions.append({
            'type': 'blank',
            'content': clean_content,
            'options': None,
            'answer': answer,
            'chapter': chapter_name
        })


def _extract_blank_answer(text):
    """从填空题文本中提取答案

    返回: (answer, clean_content)
    """
    # 模式1: 是 + 答案 + 句号
    match = re.search(r'是([^，。]+)[。，]', text)
    if match:
        answer = match.group(1).strip()
        if len(answer) > 2:  # 答案长度大于2才有效
            clean_content = text[:match.start(1)] + '____' + text[match.end(1):]
            return answer, clean_content

    # 模式2: 为 + 答案 + 句号
    match = re.search(r'为([^，。]+)[。，]', text)
    if match:
        answer = match.group(1).strip()
        if len(answer) > 2:
            clean_content = text[:match.start(1)] + '____' + text[match.end(1):]
            return answer, clean_content

    # 模式3: 转化为 + 答案 + 句号
    match = re.search(r'转化为([^，。]+)[。，]', text)
    if match:
        answer = match.group(1).strip()
        clean_content = text[:match.start(1)] + '____' + text[match.end(1):]
        return answer, clean_content

    # 模式4: 包括 + 答案 + 句号
    match = re.search(r'包括([^，。]+)[。，]', text)
    if match:
        answer = match.group(1).strip()
        clean_content = text[:match.start(1)] + '____' + text[match.end(1):]
        return answer, clean_content

    # 模式5: 即 + 答案 + 句号
    match = re.search(r'即([^，。]+)[。，]', text)
    if match:
        answer = match.group(1).strip()
        clean_content = text[:match.start(1)] + '____' + text[match.end(1):]
        return answer, clean_content

    # 模式6: 指 + 答案 + 句号
    match = re.search(r'指([^，。]+)[。，]', text)
    if match:
        answer = match.group(1).strip()
        if len(answer) > 2:
            clean_content = text[:match.start(1)] + '____' + text[match.end(1):]
            return answer, clean_content

    # 模式7: 意味着 + 答案 + 句号
    match = re.search(r'意味着([^，。]+)[。，]', text)
    if match:
        answer = match.group(1).strip()
        if len(answer) > 2:
            clean_content = text[:match.start(1)] + '____' + text[match.end(1):]
            return answer, clean_content

    # 模式8: 迎来了 + 答案 + 句号
    match = re.search(r'迎来了([^，。]+)[。，]', text)
    if match:
        answer = match.group(1).strip()
        if len(answer) > 2:
            clean_content = text[:match.start(1)] + '____' + text[match.end(1):]
            return answer, clean_content

    # 模式9: 指出 + 逗号 + 答案 + 句号
    match = re.search(r'指出[，,]([^。]+)[。，]', text)
    if match:
        answer = match.group(1).strip()
        if len(answer) > 2:
            clean_content = text[:match.start(1)] + '____' + text[match.end(1):]
            return answer, clean_content

    # 没有匹配到模式，返回空答案
    return '', text


def _parse_judges(items, questions, chapter_name, table_judge_answers):
    """解析判断题"""
    for para_idx, text, red_keys, has_red, red_full_text, _ul in items:
        m = re.match(r'^(\d+)[.、]\s*(.+)', text)
        if m:
            q_num = m.group(1)
            q_content = m.group(2).strip()
        else:
            q_num = None
            q_content = text

        # 检查内联答案
        answer = ''
        if '（√）' in q_content or '(√)' in q_content:
            answer = '√'
            q_content = q_content.replace('（√）', '').replace('(√)', '').strip()
        elif '（×）' in q_content or '(×)' in q_content:
            answer = '×'
            q_content = q_content.replace('（×）', '').replace('(×)', '').strip()

        # 如果没有内联答案，检查表格
        if not answer:
            key = q_content[:20]
            if key in table_judge_answers:
                answer = table_judge_answers[key]

        questions.append({
            'type': 'judge',
            'content': q_content,
            'options': None,
            'answer': answer,
            'chapter': chapter_name
        })


def _is_likely_stem_line(text):
    """判断是否是题干行（新格式）"""
    # 以问号、句号结尾
    if text.endswith(('？', '。', '！', '?', '.', '!')):
        return True
    # 包含"下列"、"以下"等词
    if any(word in text for word in ['下列', '以下', '以下哪', '下列哪']):
        return True
    # 包含问号
    if '？' in text or '?' in text:
        return True
    # 包含填空下划线
    if '________' in text or '______' in text:
        return True
    # 以"是"结尾且较长
    if text.endswith('是') and len(text) > 5:
        return True
    # 包含句号在中间
    if '。' in text and len(text) > 20:
        return True
    return False


def _is_option_line_v2(text, has_red):
    """判断是否是选项行（新格式：无标记选项行）

    选项行特点：
    - 短行（通常少于40字）
    - 不以句号、问号结尾
    - 不是题干特征
    """
    # 有显式选项标记的是选项行
    if re.search(r'（[A-H]）', text):
        return True
    # 短行且不以句号结尾
    if len(text) < 40 and not text.endswith(('。', '？', '！', '，', ',')):
        # 排除题干特征
        if not _is_likely_stem_line(text):
            return True
    return False


def _parse_single_choice_v2(items, questions, chapter_name):
    """解析单选题（学习通格式）

    特点：选项无标记，答案用红色字体标记整行
    处理三种选项格式：
    1. 显式标记: （A）选项 （B）选项...
    2. Tab分隔: 选项A\t（B）选项B\t（C）选项C...
    3. 单独成行: 每个选项一行（无标记），红色行=答案
    """
    current_q = None

    i = 0
    while i < len(items):
        para_idx, text, red_keys, has_red, red_full_text, _ul = items[i]

        # 检查是否有显式选项标记 （A）（B）
        has_explicit_options = bool(re.search(r'（[A-H]）', text))

        # 判断当前行是否是新题干的开始
        is_new_stem = _is_new_stem_line(text)

        # ===== 情况1: 有显式选项标记 =====
        if has_explicit_options:
            opts = _parse_options_with_first_unmarked(text)
            if current_q is not None:
                for key, opt_text in opts:
                    current_q['options'].append((key, opt_text))
                    if red_keys and key in red_keys:
                        current_q['red_keys'].append(key)
                    elif has_red:
                        if opt_text in red_full_text or red_full_text in opt_text:
                            current_q['red_keys'].append(key)
            else:
                # 没有当前题目，可能是题干+选项在同一行
                parts = re.split(r'（([A-H])）', text)
                if len(parts) >= 3:
                    stem = parts[0].strip()
                    current_q = {'stem': stem, 'options': [], 'red_keys': []}
                    for j in range(1, len(parts), 2):
                        if j + 1 < len(parts):
                            key = parts[j]
                            opt_text = parts[j + 1].strip()
                            current_q['options'].append((key, opt_text))
                            if red_keys and key in red_keys:
                                current_q['red_keys'].append(key)
                            elif has_red and opt_text in red_full_text:
                                current_q['red_keys'].append(key)
                else:
                    current_q = {'stem': text, 'options': [], 'red_keys': []}
            i += 1
            continue

        # ===== 情况2: 新题干开始 =====
        if is_new_stem:
            if current_q and current_q['stem']:
                _save_choice_question(current_q, 'single', questions, chapter_name)
            current_q = {'stem': text, 'options': [], 'red_keys': []}
            i += 1
            continue

        # ===== 情况3: 非显式选项、非新题干 → 选项行或题干续行 =====
        if current_q is None:
            current_q = {'stem': text, 'options': [], 'red_keys': []}
            i += 1
            continue

        # 如果已有4个选项，认为这是新题目
        if len(current_q['options']) >= 4:
            _save_choice_question(current_q, 'single', questions, chapter_name)
            current_q = {'stem': text, 'options': [], 'red_keys': []}
            i += 1
            continue

        # 当前题干以句号/问号结尾 → 后续行都是选项
        if current_q['stem'].endswith(('。', '？', '！')):
            _add_option_to_question(current_q, text, has_red, red_keys)
            i += 1
            continue

        # 当前题干以逗号结尾 → 题干续行
        if current_q['stem'].endswith(('，', ',')):
            current_q['stem'] += text
            i += 1
            continue

        # 关键逻辑：题干不以句号结尾
        # 收集后续行，判断是选项还是题干续行
        if not current_q['options']:
            # 还没有选项，收集后续选项行
            # 规则：连续的短行（<=40字，不以逗号结尾）视为选项
            # 遇到任何"新题干信号"停止
            option_lines = []
            j = i
            while j < len(items):
                j_text = items[j][1]
                j_has_red = items[j][3]
                j_red_keys = items[j][2]

                # 新题干信号：以句号/问号结尾且较长、有下划线、有显式选项标记
                if j > i:
                    if re.search(r'（[A-H]）', j_text):
                        break
                    if j_text.startswith('____'):
                        break
                    # 以句号/问号结尾且较长 → 新题干
                    if j_text.endswith(('。', '？', '！')) and len(j_text) > 15:
                        break
                    # 以"是"结尾且较长 → 新题干（如"主要根源是"）
                    if j_text.endswith('是') and len(j_text) > 15:
                        break

                # 判断这行是选项还是题干续行
                # 选项特征：短行（<=40字）且不以逗号结尾
                # 题干续行特征：以逗号结尾、或较长
                if len(j_text) <= 40 and not j_text.endswith(('，', ',')):
                    option_lines.append((j, j_text, j_has_red, j_red_keys))
                else:
                    # 这行不是选项（可能是题干续行或新题干），停止收集
                    break
                j += 1

            if len(option_lines) >= 2:
                # 找到多个选项行，这些是选项
                for _, opt_text, opt_has_red, opt_red_keys in option_lines:
                    _add_option_to_question(current_q, opt_text, opt_has_red, opt_red_keys)
                i = option_lines[-1][0] + 1  # 跳到最后一个选项行
                continue
            elif len(option_lines) == 1:
                # 只有一个选项行，检查它是否像选项
                _, opt_text, opt_has_red, opt_red_keys = option_lines[0]
                # 如果这行是红色的（答案标记），或者是短行，视为选项
                if opt_has_red or len(opt_text) <= 40:
                    _add_option_to_question(current_q, opt_text, opt_has_red, opt_red_keys)
                    i = option_lines[0][0] + 1
                    continue

        # 默认：题干续行
        current_q['stem'] += text
        i += 1

    # 保存最后一题
    if current_q and current_q['stem']:
        _save_choice_question(current_q, 'single', questions, chapter_name)


def _parse_multiple_choice_v2(items, questions, chapter_name):
    """解析多选题（学习通格式）

    特点：选项无标记，答案用红色字体标记整行
    """
    current_q = None

    i = 0
    while i < len(items):
        para_idx, text, red_keys, has_red, red_full_text, _ul = items[i]

        has_explicit_options = bool(re.search(r'（[A-H]）', text))
        is_new_stem = _is_new_stem_line(text)

        # ===== 情况1: 有显式选项标记 =====
        if has_explicit_options:
            opts = _parse_options_with_first_unmarked(text)
            if current_q is not None:
                for key, opt_text in opts:
                    current_q['options'].append((key, opt_text))
                    if red_keys and key in red_keys:
                        current_q['red_keys'].append(key)
                    elif has_red:
                        if opt_text in red_full_text or red_full_text in opt_text:
                            current_q['red_keys'].append(key)
            else:
                parts = re.split(r'（([A-H])）', text)
                if len(parts) >= 3:
                    stem = parts[0].strip()
                    current_q = {'stem': stem, 'options': [], 'red_keys': []}
                    for j in range(1, len(parts), 2):
                        if j + 1 < len(parts):
                            key = parts[j]
                            opt_text = parts[j + 1].strip()
                            current_q['options'].append((key, opt_text))
                            if red_keys and key in red_keys:
                                current_q['red_keys'].append(key)
                            elif has_red and opt_text in red_full_text:
                                current_q['red_keys'].append(key)
                else:
                    current_q = {'stem': text, 'options': [], 'red_keys': []}
            i += 1
            continue

        # ===== 情况2: 新题干开始 =====
        if is_new_stem:
            if current_q and current_q['stem']:
                _save_choice_question(current_q, 'multiple', questions, chapter_name)
            current_q = {'stem': text, 'options': [], 'red_keys': []}
            i += 1
            continue

        # ===== 情况3: 选项行或题干续行 =====
        if current_q is None:
            current_q = {'stem': text, 'options': [], 'red_keys': []}
            i += 1
            continue

        if len(current_q['options']) >= 4:
            _save_choice_question(current_q, 'multiple', questions, chapter_name)
            current_q = {'stem': text, 'options': [], 'red_keys': []}
            i += 1
            continue

        if current_q['stem'].endswith(('。', '？', '！')):
            _add_option_to_question(current_q, text, has_red, red_keys)
            i += 1
            continue

        if current_q['stem'].endswith(('，', ',')):
            current_q['stem'] += text
            i += 1
            continue

        # 关键逻辑：题干不以句号结尾，收集后续选项行
        if not current_q['options']:
            option_lines = []
            j = i
            while j < len(items):
                j_text = items[j][1]
                j_has_red = items[j][3]
                j_red_keys = items[j][2]

                if j > i:
                    if _is_new_stem_line(j_text) or re.search(r'（[A-H]）', j_text):
                        break
                    if j_text.startswith('____'):
                        break
                    if j_text.endswith(('。', '？', '！')) and len(j_text) > 15:
                        break
                    if j_text.endswith('是') and len(j_text) > 15:
                        break

                if len(j_text) <= 40 and not j_text.endswith(('，', ',')):
                    option_lines.append((j, j_text, j_has_red, j_red_keys))
                elif j_text.endswith(('，', ',')):
                    break
                else:
                    break
                j += 1

            if len(option_lines) >= 2:
                for _, opt_text, opt_has_red, opt_red_keys in option_lines:
                    _add_option_to_question(current_q, opt_text, opt_has_red, opt_red_keys)
                i = option_lines[-1][0] + 1
                continue
            elif len(option_lines) == 1:
                _, opt_text, opt_has_red, opt_red_keys = option_lines[0]
                if opt_has_red or len(opt_text) <= 40:
                    _add_option_to_question(current_q, opt_text, opt_has_red, opt_red_keys)
                    i = option_lines[0][0] + 1
                    continue

        current_q['stem'] += text
        i += 1

    if current_q and current_q['stem']:
        _save_choice_question(current_q, 'multiple', questions, chapter_name)


def parse_docx_to_chapters(filepath):
    """解析Word文件，返回按章节分组的题目字典

    自动检测文件格式：
    - 学习通格式（每个章节有独立题型标记）→ 使用 parse_docx_file_v2
    - 旧格式（全局题型标记）→ 使用 parse_docx_file
    """
    # 先检测是否是学习通格式
    doc = docx.Document(filepath)
    first_section_count = 0
    for p in doc.paragraphs[:300]:
        text = p.text.strip()
        if text in ('一、单项选择题', '一、单选题'):
            first_section_count += 1
    is_v2_format = first_section_count > 1  # 多个"一、单项选择题"说明是新格式

    if is_v2_format:
        questions = parse_docx_file_v2(filepath)
    else:
        questions = parse_docx_file(filepath)

    chapters = {}
    for q in questions:
        chapter_name = q.pop('chapter', '未知章节')
        if chapter_name not in chapters:
            chapters[chapter_name] = []
        chapters[chapter_name].append(q)

    return chapters


def parse_docx_file_v2(filepath):
    """解析学习通格式的Word题库文件（适配考试版）

    格式特点：每个章节有独立的题型标记（一、单项选择题 / 二、多项选择题 / 三、判断题）
    """
    doc = docx.Document(filepath)

    # 从目录中提取章节名称
    chapter_names = _extract_chapter_names_from_toc(doc)

    # 收集表格中的判断题答案
    table_judge_answers = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2:
                q_text = cells[0]
                ans_text = cells[1]
                if '（√）' in ans_text or '（×）' in ans_text:
                    answer = '√' if '√' in ans_text else '×'
                    table_judge_answers[q_text[:20]] = answer

    # 第一遍：收集所有段落并识别章节边界
    # 章节由"一、单项选择题"分隔（每次出现代表新章节）
    all_items = []
    current_type = None
    current_chapter_idx = -1
    found_first_section = False  # 是否已找到第一个题型标记

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # 跳过 [WARNING] 标记
        if '[WARNING:' in text:
            continue

        # 检测章节边界（"一、单项选择题"出现 = 新章节开始）
        if text in ('一、单项选择题', '一、单选题'):
            current_chapter_idx += 1
            current_type = 'single'
            found_first_section = True
            continue

        # 跳过第一个题型标记之前的所有内容（目录等）
        if not found_first_section:
            continue

        # 检测题型切换（在当前章节内）
        if text in ('二、多项选择题', '二、多选题'):
            current_type = 'multiple'
            continue
        elif text in ('三、判断题', '四、判断题'):
            current_type = 'judge'
            continue
        elif text in ('三、填空题',):
            current_type = 'blank'
            continue
        elif text in ('五、简答题', '六、材料论述题'):
            current_type = None
            continue

        # 跳过答案行
        if text.startswith('【答案】') or text.startswith('【参考答案】'):
            continue

        if current_type is None or current_chapter_idx < 0:
            continue

        # 获取红色标记的选项
        red_keys, has_red, red_full_text = _get_line_red_answer_keys(para)
        # 获取下划线文本（填空题答案用）
        underline_texts = _get_underline_texts(para)

        all_items.append((para_idx, current_type, text, red_keys, has_red, current_chapter_idx, red_full_text, underline_texts))

    # 第二遍：解析题目
    questions = []
    _parse_all_items_v2(all_items, questions, table_judge_answers, chapter_names)

    return questions


def _parse_all_items_v2(all_items, questions, table_judge_answers, chapter_names):
    """解析所有条目（学习通格式），按章节和题型分组"""
    # 按 (章节idx, 题型) 分组
    groups = []
    current_key = None
    current_items = []

    for item in all_items:
        para_idx, qtype, text, red_keys, has_red, chapter_idx, red_full_text, underline_texts = item
        key = (chapter_idx, qtype)

        if key != current_key:
            if current_items:
                groups.append((current_key, current_items[:]))
            current_key = key
            current_items = []

        current_items.append((para_idx, text, red_keys, has_red, red_full_text, underline_texts))

    if current_items:
        groups.append((current_key, current_items[:]))

    # 解析每组
    for (chapter_idx, qtype), items in groups:
        chapter_name = chapter_names[chapter_idx] if chapter_idx < len(chapter_names) else f'第{chapter_idx+1}章'

        if qtype == 'single':
            _parse_single_choice_v2(items, questions, chapter_name)
        elif qtype == 'multiple':
            _parse_multiple_choice_v2(items, questions, chapter_name)
        elif qtype == 'blank':
            _parse_blanks(items, questions, chapter_name)
        elif qtype == 'judge':
            _parse_judges(items, questions, chapter_name, table_judge_answers)
