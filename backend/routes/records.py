from flask import Blueprint, jsonify, request, send_file
from models import db, Question, Record, Progress, Chapter, Subject
from sqlalchemy import func
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
import io
import json
from datetime import datetime

records_bp = Blueprint('records', __name__)


@records_bp.route('/records/wrong', methods=['GET'])
def get_wrong_questions():
    """获取错题列表（所有曾经答错过的题目，按状态分为待复习和已掌握）"""
    subject_id = request.args.get('subject_id', type=int)

    # 查询所有曾经答错过至少一次的题目
    wrong_question_ids = db.session.query(Record.question_id).filter(
        Record.is_correct == False
    ).distinct().subquery()

    query = db.session.query(Question, Progress).join(
        wrong_question_ids, Question.id == wrong_question_ids.c.question_id
    ).outerjoin(Progress, Progress.question_id == Question.id)

    if subject_id:
        query = query.join(Chapter).filter(Chapter.subject_id == subject_id)

    results = query.order_by(
        func.coalesce(Progress.last_attempt_at, Question.created_at).desc()
    ).all()

    # 一次查询获取所有错题的统计信息（错误次数 + 最后错误时间）
    wrong_stats = db.session.query(
        Record.question_id,
        func.count().label('wrong_count'),
        func.max(Record.created_at).label('last_wrong_time')
    ).filter(
        Record.is_correct == False
    ).group_by(Record.question_id).all()
    stats_map = {s.question_id: s for s in wrong_stats}

    data = []
    for q, p in results:
        stat = stats_map.get(q.id)
        current_status = p.status if p else 'unanswered'
        data.append({
            'id': q.id,
            'type': q.type,
            'content': q.content,
            'chapter_id': q.chapter_id,
            'attempt_count': stat.wrong_count if stat else 0,
            'last_attempt_at': stat.last_wrong_time.strftime('%Y-%m-%d %H:%M') if stat and stat.last_wrong_time else None,
            'status': current_status,
            'mastered': current_status == 'correct'  # 标记是否已掌握
        })

    return jsonify({
        'code': 0,
        'data': data,
        'message': 'success'
    })


def get_wrong_questions_data(subject_id=None, ids=None):
    """获取错题数据的内部函数"""
    # 查询所有曾经答错过至少一次的题目
    wrong_question_ids = db.session.query(Record.question_id).filter(
        Record.is_correct == False
    ).distinct().subquery()

    query = db.session.query(Question, Progress, Chapter, Subject).join(
        wrong_question_ids, Question.id == wrong_question_ids.c.question_id
    ).outerjoin(Progress, Progress.question_id == Question.id
    ).join(Chapter, Chapter.id == Question.chapter_id
    ).join(Subject, Subject.id == Chapter.subject_id)

    if subject_id:
        query = query.filter(Chapter.subject_id == subject_id)

    # 如果指定了 ids，只查询这些题目
    if ids:
        query = query.filter(Question.id.in_(ids))

    results = query.order_by(
        func.coalesce(Progress.last_attempt_at, Question.created_at).desc()
    ).all()

    # 一次查询获取所有错题的统计信息（错误次数 + 最后错误时间）
    wrong_stats = db.session.query(
        Record.question_id,
        func.count().label('wrong_count'),
        func.max(Record.created_at).label('last_wrong_time')
    ).filter(
        Record.is_correct == False
    ).group_by(Record.question_id).all()
    stats_map = {s.question_id: s for s in wrong_stats}

    data = []
    for q, p, chapter, subject in results:
        stat = stats_map.get(q.id)
        current_status = p.status if p else 'unanswered'

        # 获取选项文本（处理 JSON 字符串或列表格式）
        options_text = ''
        if q.options:
            # 如果是字符串，先解析为 JSON
            options = q.options
            if isinstance(options, str):
                try:
                    options = json.loads(options)
                except json.JSONDecodeError:
                    options = []

            if isinstance(options, list):
                opt_lines = []
                for opt in options:
                    if isinstance(opt, dict):
                        # 支持两种选项格式：key/text 或 label/content
                        key = opt.get('key') or opt.get('label', '')
                        text = opt.get('text') or opt.get('content', '')
                        if key and text:
                            opt_lines.append(f"{key}. {text}")
                options_text = '\n'.join(opt_lines)

        data.append({
            'id': q.id,
            'subject': subject.name,
            'chapter': chapter.name,
            'type': q.type,
            'type_text': {'single': '单选题', 'multiple': '多选题', 'blank': '填空题', 'judge': '判断题'}.get(q.type, ''),
            'content': q.content,
            'options': options_text,
            'answer': q.answer,
            'explanation': q.explanation or '',
            'attempt_count': stat.wrong_count if stat else 0,
            'last_attempt_at': stat.last_wrong_time.strftime('%Y-%m-%d %H:%M') if stat and stat.last_wrong_time else '',
            'status': current_status,
            'mastered': current_status == 'correct'
        })

    return data


@records_bp.route('/export/wrong/word', methods=['GET'])
def export_wrong_to_word():
    """导出错题为 Word 文档"""
    subject_id = request.args.get('subject_id', type=int)
    tab = request.args.get('tab', 'pending')  # pending 或 mastered
    ids_str = request.args.get('ids', '')  # 逗号分隔的题目 ID

    # 解析 ids 参数
    ids = None
    if ids_str:
        try:
            ids = [int(id_str) for id_str in ids_str.split(',') if id_str.strip()]
        except ValueError:
            pass

    data = get_wrong_questions_data(subject_id, ids)

    # 根据 tab 过滤
    if tab == 'pending':
        data = [d for d in data if not d['mastered']]
        title_suffix = '待复习'
    elif tab == 'mastered':
        data = [d for d in data if d['mastered']]
        title_suffix = '已掌握'
    else:
        title_suffix = '全部'

    # 创建 Word 文档
    doc = Document()

    # 设置默认字体为宋体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(14)  # 四号字 = 14磅

    # 设置标题
    title = doc.add_heading(f'错题本 - {title_suffix}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '宋体'
        run.font.size = Pt(14)

    # 添加导出信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'导出时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}    题目数量：{len(data)} 题')
    run.font.name = '宋体'
    run.font.size = Pt(12)  # 小四号字

    doc.add_paragraph()  # 空行

    # 按科目和章节分组
    from collections import defaultdict
    grouped = defaultdict(lambda: defaultdict(list))
    for item in data:
        grouped[item['subject']][item['chapter']].append(item)

    # 写入题目
    for subject, chapters in grouped.items():
        h2 = doc.add_heading(subject, level=2)
        for run in h2.runs:
            run.font.name = '宋体'
            run.font.size = Pt(14)

        for chapter, questions in chapters.items():
            h3 = doc.add_heading(chapter, level=3)
            for run in h3.runs:
                run.font.name = '宋体'
                run.font.size = Pt(14)

            for i, q in enumerate(questions, 1):
                # 题目内容
                p = doc.add_paragraph()
                run = p.add_run(f"{i}. ")
                run.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(14)
                run = p.add_run(f"【{q['type_text']}】{q['content']}")
                run.font.name = '宋体'
                run.font.size = Pt(14)

                # 选项（如果有）
                if q['options']:
                    for opt_line in q['options'].split('\n'):
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Cm(1)
                        run = p.add_run(opt_line)
                        run.font.name = '宋体'
                        run.font.size = Pt(14)

                # 答案
                p = doc.add_paragraph()
                run = p.add_run('答案：')
                run.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(14)
                run = p.add_run(q['answer'])
                run.font.name = '宋体'
                run.font.size = Pt(14)

                # 错误信息
                p = doc.add_paragraph()
                run = p.add_run(f"错误次数：{q['attempt_count']} 次")
                run.font.name = '宋体'
                run.font.size = Pt(12)
                if q['last_attempt_at']:
                    run = p.add_run(f"    最后错误：{q['last_attempt_at']}")
                    run.font.name = '宋体'
                    run.font.size = Pt(12)

                # 解析（如果有）
                if q['explanation']:
                    p = doc.add_paragraph()
                    run = p.add_run('解析：')
                    run.bold = True
                    run.font.name = '宋体'
                    run.font.size = Pt(14)
                    run = p.add_run(q['explanation'])
                    run.font.name = '宋体'
                    run.font.size = Pt(14)

                # 分隔线
                doc.add_paragraph('─' * 50)

    # 保存到内存
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'错题本_{title_suffix}_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
    )


@records_bp.route('/export/wrong/excel', methods=['GET'])
def export_wrong_to_excel():
    """导出错题为 Excel 文件"""
    subject_id = request.args.get('subject_id', type=int)
    tab = request.args.get('tab', 'pending')  # pending 或 mastered
    ids_str = request.args.get('ids', '')  # 逗号分隔的题目 ID

    # 解析 ids 参数
    ids = None
    if ids_str:
        try:
            ids = [int(id_str) for id_str in ids_str.split(',') if id_str.strip()]
        except ValueError:
            pass

    data = get_wrong_questions_data(subject_id, ids)

    # 根据 tab 过滤
    if tab == 'pending':
        data = [d for d in data if not d['mastered']]
        title_suffix = '待复习'
    elif tab == 'mastered':
        data = [d for d in data if d['mastered']]
        title_suffix = '已掌握'
    else:
        title_suffix = '全部'

    # 创建 Excel 工作簿
    wb = Workbook()
    ws = wb.active
    ws.title = f'错题本-{title_suffix}'

    # 定义样式
    header_font = Font(name='微软雅黑', bold=True, size=12, color='FFFFFF')
    header_fill = PatternFill(start_color='FF6B9D', end_color='FF6B9D', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    content_font = Font(name='微软雅黑', size=10)
    content_alignment = Alignment(vertical='center', wrap_text=True)

    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # 设置列宽
    columns = [
        ('序号', 6),
        ('科目', 10),
        ('章节', 20),
        ('题型', 8),
        ('题目内容', 50),
        ('选项', 30),
        ('正确答案', 12),
        ('错误次数', 10),
        ('最后错误时间', 18),
        ('掌握状态', 10),
    ]

    for i, (name, width) in enumerate(columns, 1):
        ws.column_dimensions[chr(64 + i)].width = width

    # 写入表头
    for i, (name, _) in enumerate(columns, 1):
        cell = ws.cell(row=1, column=i, value=name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    # 写入数据
    for row_idx, item in enumerate(data, 2):
        values = [
            row_idx - 1,
            item['subject'],
            item['chapter'],
            item['type_text'],
            item['content'],
            item['options'],
            item['answer'],
            item['attempt_count'],
            item['last_attempt_at'],
            '已掌握' if item['mastered'] else '待复习'
        ]

        for col_idx, value in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = content_font
            cell.alignment = content_alignment
            cell.border = thin_border

        # 根据掌握状态设置行颜色
        if item['mastered']:
            for col_idx in range(1, len(columns) + 1):
                ws.cell(row=row_idx, column=col_idx).fill = PatternFill(
                    start_color='E8F5E9', end_color='E8F5E9', fill_type='solid'
                )
        else:
            for col_idx in range(1, len(columns) + 1):
                ws.cell(row=row_idx, column=col_idx).fill = PatternFill(
                    start_color='FFF3E0', end_color='FFF3E0', fill_type='solid'
                )

    # 设置行高
    for row in range(1, len(data) + 2):
        ws.row_dimensions[row].height = 30

    # 保存到内存
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f'错题本_{title_suffix}_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
    )
